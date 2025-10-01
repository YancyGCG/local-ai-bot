from __future__ import annotations

"""DOCX export helpers."""

from pathlib import Path
from typing import Callable, Dict

from docx import Document as DocumentFactory
from docx.document import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import Table
from docx.text.paragraph import Paragraph


class UnsupportedDocumentError(Exception):
    """Raised when an unsupported document type is requested."""


def export_docx(task: Dict, doc_type: str, output_path: Path, *, template_path: Path | None = None) -> Path:
    """Dispatch export to the correct formatter."""

    builders: Dict[str, Callable[[Dict, Path, Path | None], Path]] = {
        "mtl-1": _build_mtl1_docx,
        "mtl-2": _build_mtl2_docx,
        "mtl-3": _build_mtl3_docx,
    }

    try:
        builder = builders[doc_type]
    except KeyError as exc:
        raise UnsupportedDocumentError(f"Unsupported document type: {doc_type}") from exc

    return builder(task, output_path, template_path)


def _clear_document(document: Document) -> None:
    """Remove the boilerplate paragraphs from a loaded template."""

    body = document._element.body  # noqa: SLF001 - direct access required to clear body
    for element in list(body):
        if element.tag.split('}')[-1] == 'sectPr':  # preserve section properties
            continue
        body.remove(element)


def _build_document(template_path: Path | None) -> Document:
    if template_path and template_path.exists():
        doc = DocumentFactory(str(template_path))
        _clear_document(doc)
        return doc
    return DocumentFactory()


def _build_mtl1_docx(task: Dict, output_path: Path, template_path: Path | None) -> Path:
    doc = _build_document(template_path)
    meta = task.get("meta", {})

    title = _add_paragraph(doc, style="Heading 1")
    title.add_run(f"{meta.get('title', 'MTL Document')} — MTL-1")

    summary = _add_paragraph(doc, style="Normal")
    summary.add_run(
        f"Task ID: {meta.get('task_id', '')} | Version: {meta.get('version', '')} | Owner: {meta.get('owner', '')}"
    )

    details = _add_paragraph(doc, style="Normal")
    details.add_run(
        f"Updated: {meta.get('last_updated', '')} | Difficulty: {meta.get('difficulty', '')} | Est. Time: {meta.get('estimated_time_min', '')} min"
    )

    tags = ", ".join(meta.get("tags", []))
    if tags:
        _add_paragraph(doc, style="Normal").add_run(f"Tags: {tags}")

    purpose = meta.get("purpose") or task.get("purpose")
    if purpose:
        _add_heading(doc, "Purpose", level=2)
        _add_paragraph(doc, style="Normal").add_run(purpose)

    _add_heading(doc, "Tools Required", level=2)
    for tool in task.get("tools_required", []):
        text = tool.get("name", "")
        if tool.get("qty"):
            text += f" ({tool['qty']})"
        if tool.get("notes"):
            text += f" — {tool['notes']}"
        _add_bullet(doc, text)

    if task.get("prerequisites"):
        _add_heading(doc, "Prerequisites", level=2)
        for prereq in task["prerequisites"]:
            _add_bullet(doc, prereq)

    _add_heading(doc, "Success Criteria", level=2)
    for item in task.get("confirmations", []):
        accept = "; ".join(item.get("accept", []))
        _add_bullet(doc, f"{item.get('item')}: {accept}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _build_mtl2_docx(task: Dict, output_path: Path, template_path: Path | None) -> Path:
    doc = _build_document(template_path)
    meta = task.get("meta", {})

    title = _add_paragraph(doc, style="Heading 1")
    title_run = title.add_run(f"{meta.get('title', 'MTL Document')} — MTL-2")
    title_run.bold = True

    subtitle = _add_paragraph(doc, style="Normal")
    subtitle_run = subtitle.add_run(
        f"MTL Number: {meta.get('task_id', '')} | Version: {meta.get('version', '')} | "
        f"Owner: {meta.get('owner', '')} | Updated: {meta.get('last_updated', '')}"
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    difficulty_line = _add_paragraph(doc, style="Normal")
    diff_run = difficulty_line.add_run(
        f"Difficulty: {meta.get('difficulty', 'Unknown')} | Est. Time: {meta.get('estimated_time_min', '')} min"
    )
    diff_run.italic = True

    tags_line = _add_paragraph(doc, style="Normal")
    tags = ", ".join(meta.get("tags", []))
    if tags:
        tags_line.add_run(f"Tags: {tags}")

    _add_heading(doc, "Environment")
    environment = task.get("environment", {})
    _write_key_value_list(doc, "Models", environment.get("models"))
    _write_key_value_list(doc, "Software Versions", environment.get("sw_versions"))
    _write_key_value_list(doc, "Connections", environment.get("connections"))

    _add_heading(doc, "Tools Required")
    for tool in task.get("tools_required", []):
        text = tool.get("name", "")
        qty = tool.get("qty")
        if qty:
            text += f" ({qty})"
        notes = tool.get("notes")
        if notes:
            text += f" — {notes}"
        _add_bullet(doc, text)

    _add_heading(doc, "Step-by-Step")
    for step in task.get("steps", []):
        heading = _add_paragraph(doc, style="Heading 2")
        heading.add_run(f"Step {step.get('sequence')} — {step.get('id')}")

        summary = _add_paragraph(doc, style="Normal")
        if step.get("critical"):
            summary_run = summary.add_run("CRITICAL: ")
            summary_run.bold = True
            summary_run.font.highlight_color = None
        summary.add_run(step.get("text", ""))

        if step.get("warnings"):
            warn_heading = _add_paragraph(doc, style="Normal")
            warn_heading.add_run("Warnings:").bold = True
            for warning in step["warnings"]:
                _add_bullet(doc, warning)

        if step.get("tips"):
            tip_heading = _add_paragraph(doc, style="Normal")
            tip_heading.add_run("Tips:").bold = True
            for tip in step["tips"]:
                _add_bullet(doc, tip)

        if step.get("confirm"):
            confirm = _add_paragraph(doc, style="Normal")
            confirm.add_run(f"Confirm: {step['confirm']}")

        doc.add_paragraph()  # spacing between steps

    _add_heading(doc, "Final Confirmations")
    for item in task.get("confirmations", []):
        accept = "; ".join(item.get("accept", []))
        _add_bullet(doc, f"{item.get('item')}: {accept}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _write_key_value_list(doc: Document, label: str, values: list[str] | None) -> None:
    if not values:
        return
    para = _add_paragraph(doc, style="Normal")
    para.add_run(f"{label}: ")
    para.add_run(", ".join(values))


def _add_heading(doc: Document, text: str, *, level: int = 2) -> Paragraph:
    style = "Heading 1" if level == 1 else "Heading 2"
    para = _add_paragraph(doc, style=style)
    run = para.add_run(text)
    run.bold = True
    return para


def _add_bullet(doc: Document, text: str) -> Paragraph:
    try:
        return doc.add_paragraph(text, style="List Bullet")
    except KeyError:
        para = doc.add_paragraph()
        para.add_run(f"• {text}")
        return para


def _add_paragraph(doc: Document, *, style: str | None = None) -> Paragraph:
    if style is None:
        return doc.add_paragraph()
    try:
        return doc.add_paragraph(style=style)
    except KeyError:
        return doc.add_paragraph()


def _build_mtl3_docx(task: Dict, output_path: Path, template_path: Path | None) -> Path:
    doc = _build_document(template_path)
    meta = task.get("meta", {})

    title = _add_paragraph(doc, style="Heading 1")
    title.add_run(f"{meta.get('title', 'MTL Document')} — MTL-3")

    teachback = task.get("teachback", {})

    _add_heading(doc, "Teachback Prompts")
    for prompt in teachback.get("prompts", []):
        _add_bullet(doc, prompt)

    rubric = teachback.get("rubric", [])
    if rubric:
        _add_heading(doc, "Rubric", level=2)
        table = _add_table(doc, rows=1, cols=5)
        headers = ["Criterion", "Needs Work", "Good", "Better", "Best"]
        for cell, text in zip(table.rows[0].cells, headers):
            cell.text = text

        for entry in rubric:
            row = table.add_row().cells
            levels_map = entry.get("levels_map", {})
            row[0].text = entry.get("criterion", "")
            row[1].text = levels_map.get("Needs Work", "")
            row[2].text = levels_map.get("Good", "")
            row[3].text = levels_map.get("Better", "")
            row[4].text = levels_map.get("Best", "")

    _add_heading(doc, "Sign-off", level=2)
    signoff = task.get("signoff", {})
    doc.add_paragraph(f"Trainee: {signoff.get('trainee_name', '________________')}")
    doc.add_paragraph(f"Trainer: {signoff.get('trainer_name', '________________')}")
    doc.add_paragraph(f"Date: {signoff.get('date', '____-____-____')}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _add_table(doc: Document, *, rows: int, cols: int) -> Table:
    try:
        table = doc.add_table(rows=rows, cols=cols, style="Table Grid")
    except KeyError:
        table = doc.add_table(rows=rows, cols=cols)
    table.autofit = True
    return table