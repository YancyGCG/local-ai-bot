"""
Handles confidential document processing with privacy-first approach
"""


import os
import sys
import json
import click
import logging
from pathlib import Path
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, Any, Optional, List
import PyPDF2
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
import requests

from mtlgen import loader


# Data class for processed document
@dataclass
class ProcessedDocument:
    """Container for processed document data"""
    filename: str
    content: str
    metadata: Dict[str, Any]
    chunks: List[str]
    embeddings: Optional[List[List[float]]] = None
    processed_at: Optional[str] = None

# Main document processing class
class DocumentProcessor:
    def __init__(self, ollama_url: str = "http://ollama:11434"):
        self.ollama_url = ollama_url
        self.embedding_model = None
        self.setup_logging()

    def setup_logging(self):
        """Configure logging for document processing"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger("DocumentProcessor")

        def extract_mtl_task(self, processed_doc: ProcessedDocument, mtl_number: str = "MTL 2") -> Dict[str, Any]:
                """Generate structured MTL task data using the LLM and validate against the schema."""

                schema_instructions = """
Return JSON with the following structure:
{
    "meta": {
        "task_id": "MTL 2",  # exactly the provided task number (e.g., MTL 2)
        "title": "...",       # short descriptive title from the document
        "version": "GT-###",  # numbering format GT-001, GT-002, etc.
        "owner": "...",       # responsible team or person
        "last_updated": "MM-DD-YY",
        "estimated_time_min": 30,
        "difficulty": "Beginner | Intermediate | Advanced",
        "tags": ["keyword", "another keyword"],
        "purpose": "One paragraph that explains why this task matters"
    },
    "tools_required": [
        {"name": "Tool name", "qty": "1", "notes": "specific notes"}
    ],
    "prerequisites": ["Specific prerequisite"],
    "environment": {
        "models": ["Equipment models"],
        "sw_versions": ["Software versions"],
        "connections": ["Connections or interfaces"]
    },
    "steps": [
        {
            "id": "Short-Step-Identifier",
            "text": "Detailed action pulled from the document",
            "critical": true,
            "confirm": "Observable confirmation",
            "tips": ["Optional tips"],
            "warnings": ["Important warnings"],
            "screenshot": "relative/path.png"
        }
    ],
    "confirmations": [
        {"item": "What is verified", "accept": ["Acceptable outcome", "Another"]}
    ],
    "teachback": {
        "prompts": ["Coaching prompt"],
        "rubric": [
            {
                "criterion": "Skill area",
                "levels": [
                    {"label": "Needs Work", "description": "..."},
                    {"label": "Good", "description": "..."},
                    {"label": "Better", "description": "..."},
                    {"label": "Best", "description": "..."}
                ]
            }
        ]
    },
    "signoff": {
        "trainee_name": "________________",
        "trainer_name": "________________",
        "date": "____-____-____"
    }
}

Rules:
- Use ONLY data from the provided document excerpt.
- Produce at least 10 steps with explicit confirmations.
- Include at least three confirmations drawing from success criteria.
- Teachback rubric must cover four distinct criteria.
- Dates must be in MM-DD-YY. Version numbers must follow GT-###.
- Do not invent screenshots; leave empty string when not available.
- Respond with valid JSON only (no markdown, commentary, or trailing text).
"""

                prompt = (
                        "You are a technical writer creating casino gaming service documentation.\n"
                        f"Document filename: {processed_doc.filename}\n"
                        "Use the following excerpt to build an MTL pack seed.\n"
                        "Focus on accurate tool names, environment details, confirmations, and safety notes.\n\n"
                        f"Excerpt:\n{processed_doc.content[:3500]}\n\n"
                        f"Task number: {mtl_number}.\n"
                        "Follow the numbering guidance: task_id must match the task number, version uses GT-###.\n"
                        "Now generate the JSON structure described below.\n\n"
                        f"Schema Instructions:\n{schema_instructions}\n"
                )

                try:
                        health = requests.get(f"{self.ollama_url}/api/tags", timeout=5)
                        if health.status_code != 200:
                                raise RuntimeError("Ollama tags endpoint unavailable")

                        response = requests.post(
                                f"{self.ollama_url}/api/generate",
                                json={
                                        "model": "llama3",
                                        "prompt": prompt,
                                        "stream": False,
                                        "options": {
                                                "temperature": 0.2,
                                                "num_predict": 2500,
                                                "top_p": 0.9,
                                                "repeat_penalty": 1.05,
                                        },
                                },
                                timeout=180,
                        )
                        response.raise_for_status()

                        raw = response.json().get("response", "")
                        json_start = raw.find("{")
                        json_end = raw.rfind("}") + 1
                        if json_start == -1 or json_end <= json_start:
                                raise ValueError("Response did not contain JSON")

                        task_data = json.loads(raw[json_start:json_end])

                        # Ensure required collections exist
                        task_data.setdefault("meta", {})
                        meta = task_data["meta"]
                        meta.setdefault("task_id", mtl_number)
                        meta.setdefault("title", processed_doc.filename.rsplit(".", 1)[0])
                        meta.setdefault("version", "GT-001")
                        meta.setdefault("owner", "Operations")
                        meta.setdefault("last_updated", datetime.now().strftime("%m-%d-%y"))
                        meta.setdefault("estimated_time_min", 30)
                        meta.setdefault("difficulty", "Intermediate")
                        meta.setdefault("tags", ["Auto-Generated"])

                        task_data.setdefault("tools_required", [])
                        if not task_data["tools_required"]:
                                task_data["tools_required"].append({"name": "To be confirmed"})

                        task_data.setdefault("prerequisites", ["Review source document"])
                        task_data.setdefault("environment", {})
                        environment = task_data["environment"]
                        environment.setdefault("models", [])
                        environment.setdefault("sw_versions", [])
                        environment.setdefault("connections", [])

                        task_data.setdefault("steps", [])
                        if not task_data["steps"]:
                                task_data["steps"].append(
                                        {
                                                "id": "Initial-Review",
                                                "text": "Review the source procedure and confirm prerequisites are satisfied.",
                                                "critical": True,
                                                "confirm": "Prerequisites are documented and acknowledged.",
                                                "tips": [],
                                                "warnings": [],
                                                "screenshot": "",
                                        }
                                )

                        task_data.setdefault("confirmations", [])
                        if not task_data["confirmations"]:
                                task_data["confirmations"].append(
                                        {"item": "Procedure documented", "accept": ["Evidence captured in LMS"]}
                                )

                        teachback = task_data.setdefault("teachback", {})
                        teachback.setdefault("prompts", ["Walk through the full flashing procedure step-by-step."])
                        teachback.setdefault("rubric", [])
                        if not teachback["rubric"]:
                                teachback["rubric"].append(
                                        {
                                                "criterion": "Execution",
                                                "levels": [
                                                        {"label": "Needs Work", "description": "Cannot explain or perform the core steps."},
                                                        {"label": "Good", "description": "Performs with guidance but misses confirmations."},
                                                        {"label": "Better", "description": "Executes steps with minor prompts."},
                                                        {"label": "Best", "description": "Confidently executes and explains confirmations."},
                                                ],
                                        }
                                )

                        task_data.setdefault(
                                "signoff",
                                {"trainee_name": "________________", "trainer_name": "________________", "date": "____-____-____"},
                        )

                        validated = loader.validate_task_dict(task_data)
                        return validated

                except Exception as exc:  # pragma: no cover - runtime path depends on LLM
                        self.logger.error("Error extracting structured MTL task: %s", exc)
                        raise

        def extract_mtl_json(self, processed_doc: ProcessedDocument, mtl_type: str = "MTL 2") -> dict:
                """Backward-compatible wrapper returning structured task data."""

                return self.extract_mtl_task(processed_doc, mtl_number=mtl_type)
    
    def load_embedding_model(self):
        """Load local embedding model for document vectorization"""
        if self.embedding_model is None:
            try:
                # Use a lightweight model that runs locally
                self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
                self.logger.info("Loaded local embedding model")
            except Exception as e:
                self.logger.warning(f"Could not load embedding model: {e}")
                self.embedding_model = None
    
    # (Removed incomplete/duplicate extract_text_from_pdf method)
    
    def extract_text_from_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text and metadata from PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                metadata = {
                    "pages": len(pdf_reader.pages),
                    "title": pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else ''
                }
            return text.strip(), metadata
        except Exception as e:
            logging.error(f"Error processing PDF {file_path}: {e}")
            return "", {}
    
    def extract_text_from_docx(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text and metadata from DOCX"""
        try:
            doc = DocxDocument(str(file_path))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "title": ""
            }
            return text.strip(), metadata
        except Exception as e:
            self.logger.error(f"Error processing DOCX {file_path}: {e}")
            return "", {}
    
    def extract_text_from_txt(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            metadata = {
                "file_size": file_path.stat().st_size,
                "lines": len(text.split('\n'))
            }
            return text.strip(), metadata
        except Exception as e:
            self.logger.error(f"Error processing TXT {file_path}: {e}")
            return "", {}
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks for better processing"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to end at a sentence boundary
            if end < len(text):
                last_period = text.rfind('.', start, end)
                last_newline = text.rfind('\n', start, end)
                boundary = max(last_period, last_newline)
                
                if boundary > start:
                    end = boundary + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks
    
    def generate_embeddings(self, chunks: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks"""
        if self.embedding_model is None:
            self.load_embedding_model()
        
        if self.embedding_model is None:
            return []
        
        try:
            embeddings = self.embedding_model.encode(chunks)
            return embeddings.tolist()
        except Exception as e:
            self.logger.error(f"Error generating embeddings: {e}")
            return []
    
    def process_document(self, file_path: str) -> ProcessedDocument:
        """Process a document and return structured data"""
        path_obj = Path(file_path)
        
        if not path_obj.exists():
            raise FileNotFoundError(f"File not found: {path_obj}")
        
        self.logger.info(f"Processing document: {path_obj.name}")
        
        # Extract text based on file type
        file_extension = path_obj.suffix.lower()
        
        if file_extension == '.pdf':
            text, metadata = self.extract_text_from_pdf(path_obj)
        elif file_extension == '.docx':
            text, metadata = self.extract_text_from_docx(path_obj)
        elif file_extension in ['.txt', '.md']:
            text, metadata = self.extract_text_from_txt(path_obj)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        if not text:
            raise ValueError(f"No text extracted from {path_obj}")
        
        # Chunk the text
        chunks = self.chunk_text(text)
        
        # Generate embeddings
        embeddings = self.generate_embeddings(chunks)
        
        # Add processing metadata
        metadata.update({
            "file_type": file_extension,
            "file_path": str(path_obj),
            "chunks_count": len(chunks),
            "character_count": len(text),
            "word_count": len(text.split())
        })
        
        return ProcessedDocument(
            filename=path_obj.name,
            content=text,
            metadata=metadata,
            chunks=chunks,
            embeddings=embeddings
        )
    
    def analyze_document(self, processed_doc: ProcessedDocument, analysis_type: str = "summary") -> str:
        """Analyze document using Ollama"""
        try:
            if analysis_type == "summary":
                prompt = f"""Please provide a comprehensive summary of the following document:

Title: {processed_doc.filename}
Content: {processed_doc.content[:2000]}...

Provide a summary that includes:
1. Main topics and themes
2. Key findings or points
3. Overall purpose and conclusion
4. Any important details or recommendations

Summary:"""
            
            elif analysis_type == "breakdown":
                prompt = f"""Please provide a detailed breakdown and analysis of the following document:

Title: {processed_doc.filename}
Content: {processed_doc.content[:2000]}...

Provide a breakdown that includes:
1. Document structure and organization
2. Main sections and their purposes
3. Key data points, statistics, or metrics
4. Critical insights and implications
5. Potential action items or next steps

Analysis:"""
            
            elif analysis_type == "questions":
                prompt = f"""Based on the following document, generate important questions that someone should ask:

Title: {processed_doc.filename}
Content: {processed_doc.content[:2000]}...

Generate 10 thoughtful questions that would help someone better understand:
- The document's implications
- Missing information that should be investigated
- Next steps or decisions to be made
- Potential risks or opportunities

Questions:"""
            
            else:
                prompt = f"Please analyze this document: {processed_doc.content[:2000]}..."
            
            # Send request to Ollama
            response = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": "llama3",
                    "prompt": prompt,
                    "stream": False,
                    "options": {"use_mmap": False}
                },
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                return result.get('response', 'No analysis generated')
            else:
                return f"Error: Could not analyze document (Status: {response.status_code})"
                
        except Exception as e:
            self.logger.error(f"Error analyzing document: {e}")
            return f"Error analyzing document: {str(e)}"

@click.command()
@click.option('--file', '-f', required=True, help='Path to the document to process')
@click.option('--analysis', '-a', default='summary', 
              type=click.Choice(['summary', 'breakdown', 'questions']),
              help='Type of analysis to perform')
@click.option('--output', '-o', help='Output file for results (optional)')
@click.option('--verbose', '-v', is_flag=True, help='Verbose output')
def main(file: str, analysis: str, output: Optional[str], verbose: bool):
    """Process and analyze confidential documents locally"""
    
    if verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    try:
        processor = DocumentProcessor()
        
        # Process the document
        click.echo(f"📄 Processing document: {file}")
        processed_doc = processor.process_document(file)
        
        click.echo(f"✅ Document processed successfully!")
        click.echo(f"   - {processed_doc.metadata['character_count']} characters")
        click.echo(f"   - {processed_doc.metadata['word_count']} words")
        click.echo(f"   - {processed_doc.metadata['chunks_count']} chunks")
        
        # Analyze the document
        click.echo(f"\n🤖 Performing {analysis} analysis...")
        analysis_result = processor.analyze_document(processed_doc, analysis)
        
        # Output results
        result = {
            "document": {
                "filename": processed_doc.filename,
                "metadata": processed_doc.metadata
            },
            "analysis_type": analysis,
            "analysis": analysis_result,
            "processed_at": processed_doc.processed_at
        }
        
        if output:
            with open(output, 'w') as f:
                json.dump(result, f, indent=2)
            click.echo(f"📁 Results saved to: {output}")
        else:
            click.echo(f"\n📊 {analysis.title()} Analysis:")
            click.echo("=" * 50)
            click.echo(analysis_result)
            click.echo("=" * 50)
        
    except Exception as e:
        click.echo(f"❌ Error: {str(e)}", err=True)
        sys.exit(1)

if __name__ == "__main__":
    main()
