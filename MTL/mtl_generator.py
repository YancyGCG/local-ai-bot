#!/usr/bin/env python3
"""
MTL (Master Task List) Document Generator

This script reads JSON data and generates Word documents from templates,
replacing placeholders with actual data.
"""

import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


class MTLGenerator:
    def __init__(self, json_file_path):
        """Initialize the generator with JSON data."""
        self.data = self.load_json_data(json_file_path)
        
    def load_json_data(self, json_file_path):
        """Load and return JSON data from file."""
        try:
            with open(json_file_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        except FileNotFoundError:
            raise FileNotFoundError(f"JSON file not found: {json_file_path}")
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")
    
    def create_document_from_template(self, template_path=None):
        """Create a new document or load from template."""
        if template_path and os.path.exists(template_path):
            try:
                doc = Document(template_path)
                return doc
            except Exception as e:
                print(f"Warning: Could not load template {template_path}: {e}")
                print("Creating new document instead...")
        
        # Create new document if template not available
        return Document()
    
    def replace_placeholders_in_text(self, text):
        """Replace placeholders in text with actual data."""
        placeholders = {
            '{{mtl_number}}': self.data.get('mtl_number', ''),
            '{{title}}': self.data.get('title', ''),
            '{{version}}': self.data.get('version', ''),
            '{{created_date}}': self.data.get('created_date', ''),
            '{{created_by}}': self.data.get('created_by', ''),
            '{{current_date}}': datetime.now().strftime('%m/%Y'),
            '{{step_count}}': str(len(self.data.get('steps', [])))
        }
        
        for placeholder, value in placeholders.items():
            text = text.replace(placeholder, value)
        
        return text
    
    def process_document_placeholders(self, doc):
        """Process all placeholders in the document."""
        # Process paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                original_text = paragraph.text
                new_text = self.replace_placeholders_in_text(original_text)
                if new_text != original_text:
                    paragraph.text = new_text
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            original_text = paragraph.text
                            new_text = self.replace_placeholders_in_text(original_text)
                            if new_text != original_text:
                                paragraph.text = new_text
    
    def create_basic_mtl_document(self):
        """Create a basic MTL document from scratch."""
        doc = Document()
        
        # Add title
        title = doc.add_heading('MASTER TASK LIST', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document info
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('MTL Number:', self.data.get('mtl_number', '')),
            ('Title:', self.data.get('title', '')),
            ('Version:', self.data.get('version', '')),
            ('Created Date:', self.data.get('created_date', '')),
            ('Created By:', self.data.get('created_by', ''))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()  # Add space
        
        # Add steps table
        steps = self.data.get('steps', [])
        if steps:
            steps_table = doc.add_table(rows=1, cols=3)
            steps_table.style = 'Table Grid'
            
            # Header row
            hdr_cells = steps_table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'STEP'
            hdr_cells[2].text = 'TECH. INITIALS'
            
            # Add steps
            for i, step in enumerate(steps, start=1):
                row = steps_table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = step
                row.cells[2].text = ""  # Empty for technician initials
        
        return doc
    
    def generate_document(self, template_path=None, output_path=None):
        """Generate the final document."""
        if template_path:
            doc = self.create_document_from_template(template_path)
            self.process_document_placeholders(doc)
        else:
            doc = self.create_basic_mtl_document()
        
        # Determine output path
        if not output_path:
            mtl_number = self.data.get('mtl_number', 'MTL')
            output_path = f"{mtl_number}_Generated.docx"
        
        # Save document
        doc.save(output_path)
        print(f"Document generated successfully: {output_path}")
        return output_path


def main():
    """Main function to run the generator."""
    json_file = "mtl1_data.json"
    template_file = "MTL1_MasterTemplate_Placeholders.docx"
    
    try:
        # Create generator instance
        generator = MTLGenerator(json_file)
        
        # Try to use template first, fallback to basic document
        if os.path.exists(template_file):
            print(f"Using template: {template_file}")
            output_file = generator.generate_document(template_path=template_file)
        else:
            print("Template not found, creating basic document")
            output_file = generator.generate_document()
        
        print(f"✅ Successfully generated: {output_file}")
        
    except Exception as e:
        print(f"❌ Error generating document: {e}")


if __name__ == "__main__":
    main()
