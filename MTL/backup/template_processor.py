#!/usr/bin/env python3
"""
Template Analyzer and MTL Generator

This script analyzes your Word template and generates documents that match your exact format.
It can identify and work with various placeholder formats and table structures.
"""

import json
import os
import re
from docx import Document
from enhanced_mtl_generator import EnhancedMTLGenerator
import argparse


class TemplateAnalyzer:
    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = None
        
    def analyze_template(self):
        """Analyze the template structure and identify placeholders."""
        if not os.path.exists(self.template_path):
            print(f"❌ Template not found: {self.template_path}")
            return
        
        try:
            self.doc = Document(self.template_path)
            print(f"\n{'='*60}")
            print(f"TEMPLATE ANALYSIS: {os.path.basename(self.template_path)}")
            print(f"{'='*60}")
            
            self._analyze_document_structure()
            self._find_placeholders()
            self._analyze_tables()
            
        except Exception as e:
            print(f"❌ Error analyzing template: {e}")
    
    def _analyze_document_structure(self):
        """Analyze basic document structure."""
        if not self.doc:
            return
            
        print(f"📄 Document Structure:")
        print(f"   Paragraphs: {len(self.doc.paragraphs)}")
        print(f"   Tables: {len(self.doc.tables)}")
        print(f"   Sections: {len(self.doc.sections)}")
        
        # Check for images
        image_count = 0
        try:
            for rel in self.doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
        except:
            pass
        print(f"   Images: {image_count}")
    
    def _find_placeholders(self):
        """Find potential placeholders in the document."""
        if not self.doc:
            return
            
        print(f"\n🔍 Searching for Placeholders:")
        
        placeholders_found = set()
        
        # Define patterns at the start
        patterns = [
            r'\{\{[^}]+\}\}',  # {{placeholder}}
            r'\{[^}]+\}',      # {placeholder}
            r'\[[^\]]+\]',     # [placeholder]
            r'<[^>]+>',        # <placeholder>
        ]
        
        # Search in paragraphs
        for para in self.doc.paragraphs:
            text = para.text
            if text:
                for pattern in patterns:
                    matches = re.findall(pattern, text)
                    placeholders_found.update(matches)
        
        # Search in tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        if text:
                            for pattern in patterns:
                                matches = re.findall(pattern, text)
                                placeholders_found.update(matches)
        
        if placeholders_found:
            print("   Found placeholders:")
            for placeholder in sorted(placeholders_found):
                print(f"     • {placeholder}")
        else:
            print("   No obvious placeholders found (using standard format)")
            print("   Will use: {{mtl_number}}, {{title}}, {{version}}, etc.")
    
    def _analyze_tables(self):
        """Analyze table structures."""
        if not self.doc:
            return
            
        print(f"\n📊 Table Analysis:")
        
        for i, table in enumerate(self.doc.tables, 1):
            print(f"   Table {i}:")
            print(f"     Rows: {len(table.rows)}")
            print(f"     Columns: {len(table.columns)}")
            
            # Analyze first row (likely header)
            if table.rows:
                header_row = table.rows[0]
                headers = [cell.text.strip() for cell in header_row.cells]
                print(f"     Headers: {headers}")
                
                # Check if this looks like a steps table
                header_text = ' '.join(headers).lower()
                if any(word in header_text for word in ['step', '#', 'task', 'procedure', 'initial']):
                    print(f"     → Identified as STEPS TABLE")
                else:
                    print(f"     → Regular data table")


def create_custom_generator(json_file, template_file, custom_placeholders=None):
    """Create a generator with custom placeholder mapping."""
    
    class CustomMTLGenerator(EnhancedMTLGenerator):
        def _create_placeholder_mapping(self):
            # Get base mapping
            mapping = super()._create_placeholder_mapping()
            
            # Add custom placeholders if provided
            if custom_placeholders:
                mapping.update(custom_placeholders)
            
            return mapping
    
    return CustomMTLGenerator(json_file)


def main():
    parser = argparse.ArgumentParser(
        description="Analyze Word templates and generate MTL documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    # Analyze a template
    python template_processor.py --analyze MTL1_MasterTemplate_Placeholders.docx
    
    # Generate document with analysis
    python template_processor.py --json mtl1_data.json --template MTL1_MasterTemplate_Placeholders.docx
    
    # Generate with custom output name
    python template_processor.py --json mtl1_data.json --template MTL1_MasterTemplate_Placeholders.docx --output MyMTL.docx
        """
    )
    
    parser.add_argument('--analyze', help='Analyze a template file')
    parser.add_argument('--json', default='mtl1_data.json', help='JSON data file')
    parser.add_argument('--template', help='Word template file')
    parser.add_argument('--output', help='Output file name')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose output')
    
    args = parser.parse_args()
    
    # If analyze mode
    if args.analyze:
        analyzer = TemplateAnalyzer(args.analyze)
        analyzer.analyze_template()
        return
    
    # Generation mode
    if not args.template:
        print("❌ Template file required for generation. Use --template option.")
        return
    
    if not os.path.exists(args.json):
        print(f"❌ JSON file not found: {args.json}")
        return
    
    try:
        # Analyze template first if verbose
        if args.verbose:
            analyzer = TemplateAnalyzer(args.template)
            analyzer.analyze_template()
            print("\n")
        
        # Generate document
        generator = EnhancedMTLGenerator(args.json)
        output_file = generator.generate_document(
            template_path=args.template,
            output_path=args.output
        )
        
        print(f"🎉 Document generated successfully!")
        print(f"📂 Location: {os.path.abspath(output_file)}")
        
    except Exception as e:
        print(f"❌ Error: {e}")


if __name__ == "__main__":
    main()
