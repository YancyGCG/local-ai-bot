#!/usr/bin/env python3
"""
Enhanced MTL Generator with Advanced Template Processing

This version provides better support for complex Word templates with:
- Enhanced placeholder replacement in all document elements
- Better preservation of template formatting
- Support for tables, headers, footers, and text boxes
- More robust error handling for template processing
"""

import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re


class EnhancedMTLGenerator:
    def __init__(self, json_file_path):
        """Initialize the enhanced generator with JSON data."""
        self.data = self.load_json_data(json_file_path)
        self.placeholders = self._create_placeholder_mapping()
        
    def load_json_data(self, json_file_path):
        """Load and return JSON data from file."""
        try:
            with open(json_file_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        except FileNotFoundError:
            raise FileNotFoundError(f"JSON file not found: {json_file_path}")
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")
    
    def _create_placeholder_mapping(self):
        """Create a comprehensive mapping of placeholders to values."""
        current_date = datetime.now()
        steps = self.data.get('steps', [])
        
        return {
            # Basic placeholders
            '{{mtl_number}}': self.data.get('mtl_number', ''),
            '{{title}}': self.data.get('title', ''),
            '{{version}}': self.data.get('version', ''),
            '{{created_date}}': self.data.get('created_date', ''),
            '{{created_by}}': self.data.get('created_by', ''),
            
            # Generated placeholders
            '{{current_date}}': current_date.strftime('%m/%Y'),
            '{{current_date_full}}': current_date.strftime('%m/%d/%Y'),
            '{{step_count}}': str(len(steps)),
            '{{total_steps}}': str(len(steps)),
            
            # Alternative formats (case insensitive support)
            '{MTL_NUMBER}': self.data.get('mtl_number', ''),
            '{TITLE}': self.data.get('title', ''),
            '{VERSION}': self.data.get('version', ''),
            '{CREATED_DATE}': self.data.get('created_date', ''),
            '{CREATED_BY}': self.data.get('created_by', ''),
            '{CURRENT_DATE}': current_date.strftime('%m/%Y'),
            '{STEP_COUNT}': str(len(steps)),
            
            # Bracket variations
            '[mtl_number]': self.data.get('mtl_number', ''),
            '[title]': self.data.get('title', ''),
            '[version]': self.data.get('version', ''),
            '[created_date]': self.data.get('created_date', ''),
            '[created_by]': self.data.get('created_by', ''),
        }
    
    def replace_placeholders_in_text(self, text):
        """Enhanced placeholder replacement with multiple formats."""
        if not text:
            return text
            
        original_text = text
        
        # Replace all placeholders
        for placeholder, value in self.placeholders.items():
            text = text.replace(placeholder, str(value))
        
        # Log replacements for debugging
        if text != original_text:
            print(f"  Replaced: '{original_text}' → '{text}'")
        
        return text
    
    def process_paragraph_runs(self, paragraph):
        """Process individual runs within a paragraph for better placeholder replacement."""
        full_text = paragraph.text
        if not any(ph in full_text for ph in self.placeholders.keys()):
            return
        
        # Get the complete text and check for placeholders
        new_text = self.replace_placeholders_in_text(full_text)
        
        if new_text != full_text:
            # Clear existing runs and add new text
            for run in paragraph.runs:
                run.clear()
            paragraph.runs[0].text = new_text if paragraph.runs else paragraph.add_run(new_text).text
    
    def process_table_placeholders(self, table):
        """Enhanced table processing with step insertion."""
        steps = self.data.get('steps', [])
        
        # First, replace any placeholders in existing cells
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    new_text = self.replace_placeholders_in_text(original_text)
                    if new_text != original_text:
                        paragraph.text = new_text
        
        # Check if this looks like a steps table (has columns that might be for step number, description, initials)
        if len(table.columns) >= 2 and steps:
            # Look for header patterns that suggest this is a steps table
            header_row = table.rows[0]
            header_text = ' '.join([cell.text.lower() for cell in header_row.cells])
            
            if any(word in header_text for word in ['step', '#', 'task', 'procedure', 'initial']):
                print(f"  Found steps table with {len(table.rows)} existing rows")
                
                # Remove existing data rows (keep header)
                while len(table.rows) > 1:
                    table._element.remove(table.rows[-1]._element)
                
                # Add steps
                for i, step in enumerate(steps, start=1):
                    row = table.add_row()
                    row.cells[0].text = str(i)
                    row.cells[1].text = step
                    if len(row.cells) > 2:
                        row.cells[2].text = ""  # Empty for initials
                
                print(f"  Added {len(steps)} steps to table")
    
    def process_document_elements(self, doc):
        """Process all document elements for placeholders."""
        print("Processing document elements...")
        
        # Process main document paragraphs
        print("  Processing paragraphs...")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                self.process_paragraph_runs(paragraph)
        
        # Process tables
        print("  Processing tables...")
        for i, table in enumerate(doc.tables):
            print(f"    Processing table {i+1}")
            self.process_table_placeholders(table)
        
        # Process headers and footers
        print("  Processing headers and footers...")
        for section in doc.sections:
            # Process header
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        new_text = self.replace_placeholders_in_text(original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
            
            # Process footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        new_text = self.replace_placeholders_in_text(original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
    
    def create_document_from_template(self, template_path):
        """Load document from template with enhanced error handling."""
        if not template_path or not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        try:
            print(f"Loading template: {template_path}")
            doc = Document(template_path)
            print(f"Template loaded successfully with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            return doc
        except Exception as e:
            raise Exception(f"Error loading template {template_path}: {e}")
    
    def create_fallback_document(self):
        """Create a document from scratch if template is not available."""
        print("Creating fallback document...")
        doc = Document()
        
        # Add title
        title = doc.add_heading('MASTER TASK LIST', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document information
        doc.add_paragraph()
        info_para = doc.add_paragraph()
        info_para.add_run(f"MTL Number: ").bold = True
        info_para.add_run(self.data.get('mtl_number', ''))
        
        info_para = doc.add_paragraph()
        info_para.add_run(f"Title: ").bold = True
        info_para.add_run(self.data.get('title', ''))
        
        info_para = doc.add_paragraph()
        info_para.add_run(f"Version: ").bold = True
        info_para.add_run(self.data.get('version', ''))
        
        info_para = doc.add_paragraph()
        info_para.add_run(f"Created Date: ").bold = True
        info_para.add_run(self.data.get('created_date', ''))
        
        info_para = doc.add_paragraph()
        info_para.add_run(f"Created By: ").bold = True
        info_para.add_run(self.data.get('created_by', ''))
        
        doc.add_paragraph()  # Space
        
        # Add steps table
        steps = self.data.get('steps', [])
        if steps:
            steps_table = doc.add_table(rows=1, cols=3)
            steps_table.style = 'Table Grid'
            
            # Header
            hdr_cells = steps_table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'STEP'
            hdr_cells[2].text = 'TECH. INITIALS'
            
            # Make header bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add steps
            for i, step in enumerate(steps, start=1):
                row = steps_table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = step
                row.cells[2].text = ""
        
        return doc
    
    def generate_document(self, template_path=None, output_path=None):
        """Generate the final document with enhanced processing."""
        print(f"\n{'='*60}")
        print(f"ENHANCED MTL DOCUMENT GENERATION")
        print(f"{'='*60}")
        print(f"JSON Data: {self.data.get('mtl_number', 'Unknown')} - {self.data.get('title', 'Unknown')}")
        
        try:
            if template_path and os.path.exists(template_path):
                doc = self.create_document_from_template(template_path)
                self.process_document_elements(doc)
            else:
                if template_path:
                    print(f"⚠️  Template not found: {template_path}")
                print("📄 Creating fallback document")
                doc = self.create_fallback_document()
        except Exception as e:
            print(f"❌ Error processing template: {e}")
            print("📄 Creating fallback document instead")
            doc = self.create_fallback_document()
        
        # Determine output path
        if not output_path:
            mtl_number = self.data.get('mtl_number', 'MTL')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f"{mtl_number}_Enhanced_{timestamp}.docx"
        
        # Save document
        try:
            doc.save(output_path)
            print(f"\n✅ Document generated successfully!")
            print(f"📁 Output: {output_path}")
            print(f"{'='*60}\n")
            return output_path
        except Exception as e:
            raise Exception(f"Error saving document: {e}")


def main():
    """Main function to run the enhanced generator."""
    json_file = "mtl1_data.json"
    template_file = "MTL1_MasterTemplate_Placeholders.docx"
    
    try:
        # Create enhanced generator
        generator = EnhancedMTLGenerator(json_file)
        
        # Generate document using template
        output_file = generator.generate_document(template_path=template_file)
        
        print(f"🎉 Success! Your MTL document has been generated.")
        print(f"📂 File location: {os.path.abspath(output_file)}")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())
