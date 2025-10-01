#!/usr/bin/env python3
"""
Custom MTL Generator for Your Specific Template

This version is specifically designed to work with your MTL1_MasterTemplate_Placeholders.docx
template that uses {MTL_#} and {MTL_TITLE} placeholder formats.
"""

import json
import os
from docx import Document
from datetime import datetime


class CustomMTLGenerator:
    def __init__(self, json_file_path):
        """Initialize the generator with JSON data."""
        self.data = self.load_json_data(json_file_path)
        self.placeholders = self._create_placeholder_mapping()
        
    def load_json_data(self, json_file_path):
        """Load and return JSON data from file."""
        try:
            with open(json_file_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        except FileNotFoundError:
            raise FileNotFoundError(f"JSON file not found: {json_file_path}")
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")
    
    def _create_placeholder_mapping(self):
        """Create placeholder mapping specifically for your template format."""
        current_date = datetime.now()
        steps = self.data.get('steps', [])
        
        return {
            # Your template's specific format
            '{MTL_#}': self.data.get('mtl_number', ''),
            '{MTL_TITLE}': self.data.get('title', ''),
            '{VERSION}': self.data.get('version', ''),
            '{CREATED_DATE}': self.data.get('created_date', ''),
            '{CREATED_BY}': self.data.get('created_by', ''),
            '{CURRENT_DATE}': current_date.strftime('%m/%Y'),
            '{STEP_COUNT}': str(len(steps)),
            
            # Also support common variations
            '{{mtl_number}}': self.data.get('mtl_number', ''),
            '{{title}}': self.data.get('title', ''),
            '{{version}}': self.data.get('version', ''),
            '{{created_date}}': self.data.get('created_date', ''),
            '{{created_by}}': self.data.get('created_by', ''),
            '{{current_date}}': current_date.strftime('%m/%Y'),
            '{{step_count}}': str(len(steps)),
            
            # Uppercase versions
            '{MTL_NUMBER}': self.data.get('mtl_number', ''),
            '{TITLE}': self.data.get('title', ''),
        }
    
    def replace_placeholders_in_text(self, text):
        """Replace placeholders in text with actual data."""
        if not text:
            return text
            
        original_text = text
        
        # Replace all placeholders
        for placeholder, value in self.placeholders.items():
            text = text.replace(placeholder, str(value))
        
        # Log replacements for debugging
        if text != original_text:
            print(f"  Replaced: '{original_text}' → '{text}'")
        
        return text
    
    def process_all_text_elements(self, doc):
        """Process all text elements in the document."""
        print("Processing all text elements...")
        
        # Process paragraphs
        print("  Processing paragraphs...")
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                original_text = paragraph.text
                new_text = self.replace_placeholders_in_text(original_text)
                if new_text != original_text:
                    # Clear and replace text while preserving formatting
                    for run in paragraph.runs:
                        run.text = ""
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.add_run(new_text)
        
        # Process tables
        print("  Processing tables...")
        for i, table in enumerate(doc.tables):
            print(f"    Processing table {i+1}")
            self.process_table(table, i+1)
    
    def process_table(self, table, table_num):
        """Process a specific table."""
        steps = self.data.get('steps', [])
        
        # Process all existing text in the table for placeholders
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        new_text = self.replace_placeholders_in_text(original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
        
        # Check if this is the steps table (Table 1 based on analysis)
        if table_num == 1 and len(table.columns) >= 3 and steps:
            print(f"    Identified as steps table - processing {len(steps)} steps")
            
            # Keep only the header row (first row) and remove others
            while len(table.rows) > 1:
                table._element.remove(table.rows[-1]._element)
            
            # Add steps rows
            for i, step in enumerate(steps, start=1):
                row = table.add_row()
                # Set step number in first column
                row.cells[0].text = str(i)
                # Set step description in second column
                row.cells[1].text = step
                # Leave third column empty (for initials)
                row.cells[2].text = ""
            
            print(f"    Added {len(steps)} steps to table")
    
    def generate_document(self, template_path, output_path=None):
        """Generate the final document."""
        print(f"\n{'='*60}")
        print(f"CUSTOM MTL DOCUMENT GENERATION")
        print(f"{'='*60}")
        print(f"Data: {self.data.get('mtl_number', 'Unknown')} - {self.data.get('title', 'Unknown')}")
        print(f"Template: {os.path.basename(template_path)}")
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        try:
            # Load template
            print(f"Loading template...")
            doc = Document(template_path)
            print(f"Template loaded successfully")
            
            # Process all content
            self.process_all_text_elements(doc)
            
            # Determine output path
            if not output_path:
                mtl_number = self.data.get('mtl_number', 'MTL')
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = f"{mtl_number}_Custom_{timestamp}.docx"
            
            # Save document
            doc.save(output_path)
            print(f"\n✅ Document generated successfully!")
            print(f"📁 Output: {output_path}")
            print(f"{'='*60}\n")
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error processing template: {e}")


def main():
    """Main function to run the custom generator."""
    json_file = "mtl1_data.json"
    template_file = "MTL1_MasterTemplate_Placeholders.docx"
    
    try:
        # Create custom generator
        generator = CustomMTLGenerator(json_file)
        
        # Generate document using your specific template
        output_file = generator.generate_document(template_path=template_file)
        
        print(f"🎉 Success! Your custom MTL document has been generated.")
        print(f"📂 File location: {os.path.abspath(output_file)}")
        
        # Also test with network data
        print(f"\nGenerating network configuration document...")
        network_generator = CustomMTLGenerator("example_network_data.json")
        network_output = network_generator.generate_document(template_path=template_file)
        print(f"📂 Network document: {os.path.abspath(network_output)}")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())
