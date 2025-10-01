#!/usr/bin/env python3
"""
Final MTL Generator - Complete Template Match

This version implements all the graphical elements and formatting to exactly match
your MTL template including blue bars, proper headers, footers, and detailed formatting.
"""

import json
import os
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


class FinalMTLGenerator:
    def __init__(self, json_file_path):
        """Initialize the generator with JSON data."""
        self.data = self.load_json_data(json_file_path)
        self.placeholders = self._create_placeholder_mapping()
        
    def load_json_data(self, json_file_path):
        """Load and return JSON data from file."""
        try:
            with open(json_file_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        except FileNotFoundError:
            raise FileNotFoundError(f"JSON file not found: {json_file_path}")
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")
    
    def _create_placeholder_mapping(self):
        """Create placeholder mapping for both old and new JSON structures."""
        current_date = datetime.now()
        steps = self.data.get('STEPS', [])
        
        # Support both MTL_# and MTL_NUMBER for backward compatibility
        mtl_number = self.data.get('MTL_NUMBER', self.data.get('MTL_#', ''))
        
        return {
            # Core required fields
            '{MTL_NUMBER}': mtl_number,
            '{MTL_TITLE}': self.data.get('MTL_TITLE', ''),
            '{VERSION_NUMBER}': self.data.get('VERSION_NUMBER', ''),
            '{REVISION_NUMBER}': self.data.get('REVISION_NUMBER', ''),
            '{CREATED_DATE}': self.data.get('CREATED_DATE', ''),
            '{CREATED_BY}': self.data.get('CREATED_BY', ''),
            '{CURRENT_DATE}': current_date.strftime('%m/%Y'),
            '{STEP_COUNT}': str(len(steps)),
            
            # Common date placeholder variations
            'XX-XX-XXXX': self.data.get('CREATED_DATE', ''),
            'XX/XX/XX': self.data.get('CREATED_DATE', ''),
            '{DATE}': self.data.get('CREATED_DATE', ''),
            '{REVIEW_DATE}': self.data.get('CREATED_DATE', ''),
            
            # New optional fields
            '{CATEGORY}': self.data.get('CATEGORY', ''),
            '{ESTIMATED_TIME}': self.data.get('ESTIMATED_TIME', ''),
            
            # Legacy support for old format with # symbol
            '{MTL_#}': mtl_number,
            
            # Alternative formats for compatibility
            '{{MTL_NUMBER}}': mtl_number,
            '{{MTL_TITLE}}': self.data.get('MTL_TITLE', ''),
            '{{VERSION_NUMBER}}': self.data.get('VERSION_NUMBER', ''),
            '{{REVISION_NUMBER}}': self.data.get('REVISION_NUMBER', ''),
            '{{CREATED_DATE}}': self.data.get('CREATED_DATE', ''),
            '{{CREATED_BY}}': self.data.get('CREATED_BY', ''),
            '{{CATEGORY}}': self.data.get('CATEGORY', ''),
            '{{ESTIMATED_TIME}}': self.data.get('ESTIMATED_TIME', ''),
            
            # Double-bracket legacy support
            '{{MTL_#}}': mtl_number,
            
            # Clean versions without brackets
            'MTL_NUMBER': mtl_number,
            'MTL_TITLE': self.data.get('MTL_TITLE', ''),
            'MTL_#': mtl_number,  # Legacy support
            'CATEGORY': self.data.get('CATEGORY', ''),
            'ESTIMATED_TIME': self.data.get('ESTIMATED_TIME', ''),
        }
    
    def replace_placeholders_in_text(self, text):
        """Replace placeholders in text with actual data."""
        if not text:
            return text
            
        original_text = text
        
        # Replace all placeholders
        for placeholder, value in self.placeholders.items():
            text = text.replace(placeholder, str(value))
        
        # Log replacements for debugging
        if text != original_text:
            print(f"  Replaced: '{original_text}' → '{text}'")
        
        return text
    
    def add_blue_shading_to_cell(self, cell):
        """Add blue shading to a table cell."""
        try:
            # Get the cell's paragraph
            paragraph = cell.paragraphs[0]
            
            # Access the cell properties
            cell_xml = cell._tc
            cell_pr = cell_xml.get_or_add_tcPr()
            
            # Create shading element
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), '4472C4')  # Blue color
            
            cell_pr.append(shading)
            
            # Make text white and bold
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                run.font.bold = True
        except Exception as e:
            print(f"  Warning: Could not apply blue shading: {e}")
    
    def process_template_document(self, doc):
        """Process the template document with enhanced formatting."""
        print("Processing template document with enhanced formatting...")
        
        # Process all paragraphs first
        print("  Processing paragraphs...")
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                original_text = paragraph.text
                new_text = self.replace_placeholders_in_text(original_text)
                if new_text != original_text:
                    # Replace text while preserving/adding formatting
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                        # Apply bold formatting if this is the MTL_TITLE
                        if '{MTL_TITLE}' in original_text:
                            paragraph.runs[0].font.bold = True
                        # Clear other runs
                        for run in paragraph.runs[1:]:
                            run.clear()
                    else:
                        run = paragraph.add_run(new_text)
                        # Apply bold formatting if this is the MTL_TITLE
                        if '{MTL_TITLE}' in original_text:
                            run.font.bold = True
        
        # Process tables with enhanced formatting
        print("  Processing tables...")
        for i, table in enumerate(doc.tables):
            print(f"    Processing table {i+1}")
            self.process_enhanced_table(table, i+1)
        
        # Process headers and footers
        print("  Processing headers and footers...")
        self.process_headers_and_footers(doc)
    
    def process_enhanced_table(self, table, table_num):
        """Process table with enhanced formatting including blue bars."""
        steps = self.data.get('STEPS', [])
        
        # First, replace any existing placeholders in the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        new_text = self.replace_placeholders_in_text(original_text)
                        if new_text != original_text:
                            self.replace_text_with_formatting(paragraph, original_text, new_text)
        
        # Check if this is the main steps table
        if table_num == 1 and len(table.columns) >= 3 and steps:
            print(f"    Processing main steps table with {len(steps)} steps")
            
            # Find the header row with "Column 1", "Column 2", "Column 3" and apply blue shading
            for row_idx, row in enumerate(table.rows):
                row_text = ' '.join([cell.text.strip() for cell in row.cells]).lower()
                if 'column 1' in row_text or 'column' in row_text:
                    print(f"    Found column header row {row_idx}, applying blue shading")
                    for cell in row.cells:
                        self.add_blue_shading_to_cell(cell)
                    break
            
            # Find the steps header row with "#", "STEP", "TECH. INITIALS"
            steps_header_row = None
            for row_idx, row in enumerate(table.rows):
                row_text = ' '.join([cell.text.strip() for cell in row.cells]).lower()
                if '#' in row_text and ('step' in row_text or 'tech' in row_text):
                    steps_header_row = row_idx
                    print(f"    Found steps header at row {row_idx}")
                    
                    # Enhance the STEP header with detailed description
                    if len(row.cells) >= 2:
                        step_cell = row.cells[1]
                        # Clear existing content
                        step_cell.text = ""
                        
                        # Add STEP header
                        step_para = step_cell.paragraphs[0]
                        step_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the STEP header
                        step_run = step_para.add_run("STEP")
                        step_run.font.bold = True
                        step_run.font.size = Pt(11)
                        
                        # Add detailed description in smaller text
                        detail_para = step_cell.add_paragraph()
                        detail_text = ("(Provide limited detail for each step without being too wordy. "
                                     "This is a checklist for a trained, experienced employee. "
                                     "For more detail, see the MTL2 for this task. Avoid inserting "
                                     "photos or other attachments on MTL1.)")
                        detail_run = detail_para.add_run(detail_text)
                        detail_run.font.size = Pt(8)
                        detail_run.font.italic = True
                    break
            
            # Remove existing step rows (keep headers)
            if steps_header_row is not None:
                # Remove rows after the steps header
                rows_to_remove = []
                for row_idx in range(steps_header_row + 1, len(table.rows)):
                    rows_to_remove.append(row_idx)
                
                # Remove from end to beginning to maintain indices
                for row_idx in reversed(rows_to_remove):
                    table._element.remove(table.rows[row_idx]._element)
                
                # Add new step rows
                for i, step in enumerate(steps, start=1):
                    row = table.add_row()
                    row.cells[0].text = str(i)
                    row.cells[1].text = step
                    if len(row.cells) > 2:
                        row.cells[2].text = ""  # Empty for initials
                
                print(f"    Added {len(steps)} step rows")
    
    def process_headers_and_footers(self, doc):
        """Process headers and footers with placeholder replacement."""
        for section in doc.sections:
            # Process header
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        new_text = self.replace_placeholders_in_text(original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
            
            # Process footer - handle both paragraphs and tables
            if section.footer:
                # Process footer paragraphs
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        new_text = self.replace_placeholders_in_text(original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
                
                # Process footer tables (for the footer data table)
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph.text.strip():
                                    original_text = paragraph.text
                                    new_text = self.replace_placeholders_in_text(original_text)
                                    if new_text != original_text:
                                        paragraph.text = new_text
                                        print(f"  Replaced: '{original_text}' → '{new_text}'")
    
    # def add_default_footer(self, footer):
    #     """Add default footer content with placeholders - DISABLED to preserve template footer format."""
    #     pass
    
    def create_fallback_document(self):
        """Create a comprehensive fallback document with all formatting."""
        print("Creating enhanced fallback document...")
        doc = Document()
        
        # Add main title
        title = doc.add_heading('MASTER TASK LIST', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add MTL_TITLE left-justified
        mtl_title_para = doc.add_paragraph()
        mtl_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        mtl_title_run = mtl_title_para.add_run(self.data.get('MTL_TITLE', ''))
        mtl_title_run.font.bold = True
        mtl_title_run.font.size = Pt(14)
        
        doc.add_paragraph()  # Space
        
        # Create main table with proper structure
        steps = self.data.get('STEPS', [])
        if steps:
            # Create table with blue header
            main_table = doc.add_table(rows=3, cols=3)
            main_table.style = 'Table Grid'
            main_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Blue bar with Column headers
            col_row = main_table.rows[0]
            col_row.cells[0].text = "Column 1"
            col_row.cells[1].text = "Column 2"
            col_row.cells[2].text = "Column 3"
            
            for cell in col_row.cells:
                self.add_blue_shading_to_cell(cell)
            
            # Steps header row
            header_row = main_table.rows[1]
            header_row.cells[0].text = "#"
            
            # Enhanced STEP header
            step_cell = header_row.cells[1]
            step_cell.text = ""
            step_para = step_cell.paragraphs[0]
            step_run = step_para.add_run("STEP")
            step_run.font.bold = True
            
            detail_para = step_cell.add_paragraph()
            detail_text = ("(Provide limited detail for each step without being too wordy. "
                         "This is a checklist for a trained, experienced employee. "
                         "For more detail, see the MTL2 for this task. Avoid inserting "
                         "photos or other attachments on MTL1.)")
            detail_run = detail_para.add_run(detail_text)
            detail_run.font.size = Pt(8)
            detail_run.font.italic = True
            
            header_row.cells[2].text = "TECH. INITIALS"
            
            # Make header row bold
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not run.font.italic:  # Don't make the detail text bold
                            run.font.bold = True
            
            # Remove the empty third row
            main_table._element.remove(main_table.rows[2]._element)
            
            # Add step rows
            for i, step in enumerate(steps, start=1):
                row = main_table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = step
                row.cells[2].text = ""
        
        # Footer will be handled by template or left empty to preserve original format
        
        return doc
    
    def generate_document(self, template_path=None, output_path=None):
        """Generate the final document with full formatting."""
        print(f"\n{'='*60}")
        print(f"FINAL MTL DOCUMENT GENERATION")
        print(f"{'='*60}")
        print(f"Data: {self.data.get('MTL_NUMBER', self.data.get('MTL_#', 'Unknown'))} - {self.data.get('MTL_TITLE', 'Unknown')}")
        
        try:
            if template_path and os.path.exists(template_path):
                print(f"Loading template: {os.path.basename(template_path)}")
                doc = Document(template_path)
                self.process_template_document(doc)
            else:
                if template_path:
                    print(f"⚠️  Template not found: {template_path}")
                print("📄 Creating enhanced fallback document")
                doc = self.create_fallback_document()
        except Exception as e:
            print(f"❌ Error processing template: {e}")
            print("📄 Creating fallback document instead")
            doc = self.create_fallback_document()
        
        # Determine output path
        if not output_path:
            mtl_number = self.data.get('MTL_NUMBER', self.data.get('MTL_#', 'MTL'))
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f"{mtl_number}_Final_{timestamp}.docx"
        
        # Save document
        try:
            doc.save(output_path)
            print(f"\n✅ Final document generated successfully!")
            print(f"📁 Output: {output_path}")
            print(f"✨ Features included:")
            print(f"   - MASTER TASK LIST header")
            print(f"   - Left-justified MTL_TITLE")
            print(f"   - Blue column headers")
            print(f"   - Enhanced STEP description")
            print(f"   - Template footer format preserved")
            print(f"   - {len(self.data.get('STEPS', []))} steps populated")
            print(f"{'='*60}\n")
            return output_path
        except Exception as e:
            raise Exception(f"Error saving document: {e}")
    
    def replace_text_with_formatting(self, paragraph, original_text, new_text):
        """Replace text in paragraph and apply bold formatting if it's MTL_TITLE."""
        if '{MTL_TITLE}' in original_text:
            # Clear the paragraph and add new text with bold formatting
            paragraph.clear()
            run = paragraph.add_run(new_text)
            run.bold = True
        else:
            # Normal text replacement
            paragraph.text = new_text


def main():
    """Main function to run the final generator."""
    import sys
    
    # Check for command line argument
    if len(sys.argv) > 1:
        json_file = sys.argv[1]
    else:
        json_file = "mtl1_data.json"
    
    template_file = "MTL1_MasterTemplate_Placeholders.docx"
    
    try:
        # Check if JSON file exists
        if not os.path.exists(json_file):
            print(f"❌ Error: JSON file '{json_file}' not found.")
            return 1
            
        # Check if template exists in current directory or Docs subdirectory
        template_paths = [
            template_file,
            f"Docs/{template_file}",
            f"../Docs/{template_file}"
        ]
        
        template_found = None
        for path in template_paths:
            if os.path.exists(path):
                template_found = path
                break
        
        # Create final generator
        generator = FinalMTLGenerator(json_file)
        
        # Generate document
        output_file = generator.generate_document(template_path=template_found)
        
        print(f"🎉 SUCCESS! Your final MTL document has been generated.")
        print(f"📂 File location: {os.path.abspath(output_file)}")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())
