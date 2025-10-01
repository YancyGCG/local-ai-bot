from docx import Document

doc = Document()
doc.add_heading('MASTER TASK LIST', 0)
doc.add_paragraph('MTL 2')
doc.add_paragraph('JCM GEN5 Printer Flashing')

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '#'
hdr_cells[1].text = 'STEP'
hdr_cells[2].text = 'TECH. INITIALS'

steps = [
    "Identify GEN5 printer model and firmware version",
    "Power ON printer and load tickets",
    "Connect PC to printer via USB",
    "Launch JCM Device Firmware Upgrade Downloader (FLDFU)",
    "Select and verify correct firmware file",
    "Click “Full Upgrade” (ensure “Erase User Settings” is selected)",
    "Monitor progress, wait for completion message",
    "Print configuration ticket to verify firmware version",
    "Test printer operation with host machine",
    "Update records and close work order"
]
for i, step in enumerate(steps, start=1):
    row = table.add_row()
    row.cells[0].text = str(i)
    row.cells[1].text = step
    row.cells[2].text = ""

doc.save("MTL1_Filled_Minimal.docx")
