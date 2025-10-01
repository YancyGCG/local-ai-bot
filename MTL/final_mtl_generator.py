#!/usr/bin/env python3
"""
Final MTL Generator - Complete Template Match

This version implements all the graphical elements and formatting to exactly match
your MTL template including blue bars, proper headers, footers, and detailed formatting.
"""

import json
import os
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


class FinalMTLGenerator:
    def __init__(self, json_file_path):
        """Initialize the generator with JSON data."""
        self.data = self.load_json_data(json_file_path)
        self.placeholders = self._create_placeholder_mapping()
        
    def load_json_data(self, json_file_path):
        """Load and return JSON data from file."""
        try:
            with open(json_file_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        except FileNotFoundError:
            raise FileNotFoundError(f"JSON file not found: {json_file_path}")
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")
    
    def _create_placeholder_mapping(self):
        """Create placeholder mapping using your exact template placeholders."""
        current_date = datetime.now()
        steps = self.data.get('STEPS', [])
        
        return {
            # Your exact placeholders
            '{MTL_TITLE}': self.data.get('MTL_TITLE', ''),
            '{VERSION}': self.data.get('VERSION_NUMBER', self.data.get('VERSION', '')),
            '{REV}': self.data.get('REVISION_NUMBER', '1'),
            '{TODAYS_DATE}': current_date.strftime('%m-%d-%Y'),
            '{CREATED_BY}': self.data.get('CREATED_BY', ''),
            '{PRE_REQ}': '\n'.join(self.data.get('PRE_REQS', [])),
            '{REQ_EQUIP}': '\n'.join(self.data.get('EQUIPMENT_LIST', [])),
            '{COMP_CRIT}': '\n'.join(self.data.get('COMPLETION_CRITERIA', [])),
            '{CURRENT_DATE}': current_date.strftime('%m/%Y'),
            '{STEP_COUNT}': str(len(steps)),
            
            # Additional common variations for compatibility
            '{MTL_NUMBER}': self.data.get('MTL_NUMBER', self.data.get('MTL_#', '')),
            '{VERSION_NUMBER}': self.data.get('VERSION_NUMBER', self.data.get('VERSION', '')),
            '{REVISION_NUMBER}': self.data.get('REVISION_NUMBER', '1'),
            '{CREATED_DATE}': current_date.strftime('%m-%d-%Y'),
        }
    
    def replace_placeholders_in_text(self, text):
        """Replace placeholders in text with actual data."""
        if not text:
            return text
            
        original_text = text
    
        # Replace all placeholders
        for placeholder, value in self.placeholders.items():
            text = text.replace(placeholder, str(value))
    
        # Log replacements for debugging
        if text != original_text:
            print(f"  Replaced: '{original_text}' → '{text}'")
    
        return text
    
    def add_blue_shading_to_cell(self, cell):
        """Add blue shading to a table cell with improved method."""
        try:
            # Access the cell properties
            cell_xml = cell._tc
            cell_pr = cell_xml.get_or_add_tcPr()
            
            # Create shading element with new blue color #00699b
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), '00699B')  # Updated blue color #00699b
            
            cell_pr.append(shading)
            
            # Make text white and bold for all paragraphs in the cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                    run.font.bold = True
                # If no runs exist, create one and format it
                if not paragraph.runs and paragraph.text:
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
            
            print(f"    Applied blue shading to cell: {cell.text}")
        except Exception as e:
            print(f"  Warning: Could not apply blue shading: {e}")
    
    def adjust_table_width(self, table):
        """Adjust table width to fit within 0.5 inch margins."""
        try:
            # Set table width to fill available space (8.5" - 1" for margins = 7.5")
            table.width = Inches(7.5)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            print(f"    Adjusted table width to 7.5 inches")
        except Exception as e:
            print(f"  Warning: Could not adjust table width: {e}")
    
    def set_document_margins(self, doc):
        """Set document margins to 0.5 inches."""
        try:
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5) 
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            print("  Set document margins to 0.5 inches")
        except Exception as e:
            print(f"  Warning: Could not set margins: {e}")
    
    def process_template_document(self, doc):
        """Process the template document with enhanced formatting."""
        print("Processing template document with enhanced formatting...")
        
        # Set document margins first
        self.set_document_margins(doc)
        
        # Determine if this is an MTL2 template
        is_mtl2 = False
        mtl_number = self.data.get('MTL_NUMBER', '')
        if 'MTL 2' in mtl_number or 'MTL2' in mtl_number:
            is_mtl2 = True
            print("  Detected MTL2 template, applying specialized formatting")
        
        # Process all paragraphs first
        print("  Processing paragraphs...")
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Try to apply special formatting for bullet points and checkboxes
                if not self.process_placeholders_with_formatting(paragraph):
                    # If no special formatting was needed, just replace placeholders
                    original_text = paragraph.text
                    new_text = self.replace_placeholders_in_text(original_text)
                    if new_text != original_text:
                        # Replace text while preserving/adding formatting
                        if paragraph.runs:
                            paragraph.runs[0].text = new_text
                            # Apply bold formatting if this is the MTL_TITLE
                            if '{MTL_TITLE}' in original_text:
                                paragraph.runs[0].font.bold = True
                            # Clear other runs
                            for run in paragraph.runs[1:]:
                                run.clear()
                        else:
                            run = paragraph.add_run(new_text)
                            # Apply bold formatting if this is the MTL_TITLE
                            if '{MTL_TITLE}' in original_text:
                                run.font.bold = True
    
        # Process tables with enhanced formatting
        print("  Processing tables...")
        for i, table in enumerate(doc.tables):
            print(f"    Processing table {i+1}")
            self.process_enhanced_table(table, i+1, is_mtl2)
        
        # Process headers and footers
        print("  Processing headers and footers...")
        self.process_headers_and_footers(doc)
        
        # Add enhanced footer section for MTL 2 & 3
        self.add_enhanced_footer_section(doc)
    
    def process_enhanced_table(self, table, table_num, is_mtl2=False):
        """Process table with enhanced formatting including blue bars."""
        steps = self.data.get('STEPS', [])
        
        # Adjust table width for 0.5" margins
        self.adjust_table_width(table)
        
        # First, replace any existing placeholders in the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        # Check for special formatting needs (bullet points, checkboxes)
                        if not self.process_placeholders_with_formatting(paragraph):
                            # Regular placeholder replacement
                            original_text = paragraph.text
                            new_text = self.replace_placeholders_in_text(original_text)
                            if new_text != original_text:
                                self.replace_text_with_formatting(paragraph, original_text, new_text)
        
        # Check if this is the main steps table
        if table_num == 1 and len(table.columns) >= 3 and steps:
            print(f"    Processing main steps table with {len(steps)} steps")
            
            # Apply blue shading to the correct header row based on your reference images
            # The blue row should contain "MASTER TASK LIST" and "MTL X" 
            for row_idx, row in enumerate(table.rows):
                row_text = ' '.join([cell.text.strip() for cell in row.cells])
                print(f"    Row {row_idx} text: '{row_text[:50]}...'")
                
                # Apply blue shading to the row containing "MASTER TASK LIST" 
                if 'MASTER TASK LIST' in row_text.upper():
                    print(f"    Applying blue shading to MASTER TASK LIST header row {row_idx}")
                    for cell_idx, cell in enumerate(row.cells):
                        self.add_blue_shading_to_cell(cell)
                    break  # Only shade one row
            
            # Find the steps header row with "#", "STEP", "TECH. INITIALS"
            steps_header_row = None
            for row_idx, row in enumerate(table.rows):
                row_text = ' '.join([cell.text.strip() for cell in row.cells]).lower()
                if '#' in row_text and ('step' in row_text or 'tech' in row_text):
                    steps_header_row = row_idx
                    print(f"    Found steps header at row {row_idx}")
                    
                    # Enhance the STEP header with detailed description
                    if len(row.cells) >= 2:
                        step_cell = row.cells[1]
                        # Clear existing content
                        step_cell.text = ""
                        
                        # Add STEP header
                        step_para = step_cell.paragraphs[0]
                        step_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the STEP header
                        step_run = step_para.add_run("STEP")
                        step_run.font.bold = True
                        step_run.font.size = Pt(11)
                        
                        # For MTL2, include additional columns formatting
                        if is_mtl2 and len(row.cells) > 2:
                            # Set MTL2 specific column headers if they're empty
                            col_headers = ["START DATE", "STOP DATE", "TRAINEE INITIALS", "TRAINER INITIALS"]
                            for i, header in enumerate(col_headers, 2):
                                if i < len(row.cells) and not row.cells[i].text.strip():
                                    header_para = row.cells[i].paragraphs[0] 
                                    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    header_run = header_para.add_run(header)
                                    header_run.font.bold = True
                                    header_run.font.size = Pt(10)
                        
                        # Add detailed description in smaller text for MTL1
                        if not is_mtl2:
                            detail_para = step_cell.add_paragraph()
                            detail_text = ("(Provide limited detail for each step without being too wordy. "
                                         "This is a checklist for a trained, experienced employee. "
                                         "For more detail, see the MTL2 for this task. Avoid inserting "
                                         "photos or other attachments on MTL1.)")
                            detail_run = detail_para.add_run(detail_text)
                            detail_run.font.size = Pt(8)
                            detail_run.font.italic = True
                    break
            
            # Remove existing step rows (keep headers)
            if steps_header_row is not None:
                # Remove rows after the steps header
                rows_to_remove = []
                for row_idx in range(steps_header_row + 1, len(table.rows)):
                    rows_to_remove.append(row_idx)
                
                # Remove from end to beginning to maintain indices
                for row_idx in reversed(rows_to_remove):
                    table._element.remove(table.rows[row_idx]._element)
                
                # Add new step rows
                for i, step in enumerate(steps, start=1):
                    row = table.add_row()
                    row.cells[0].text = str(i)
                    row.cells[1].text = step
                    
                    # For MTL2, make sure we have enough columns
                    if is_mtl2:
                        # Ensure cells are empty for trainee completion
                        for j in range(2, min(len(row.cells), 6)):
                            if j < len(row.cells):
                                row.cells[j].text = ""
                    else:
                        # Standard MTL1 format
                        if len(row.cells) > 2:
                            row.cells[2].text = ""  # Empty for initials
            
            print(f"    Added {len(steps)} step rows")
    
    def process_headers_and_footers(self, doc):
        """Process headers and footers with placeholder replacement."""
        for section in doc.sections:
            # Process header
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        new_text = self.replace_placeholders_in_text(original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
            
            # Process footer - handle both paragraphs and tables
            if section.footer:
                # Process footer paragraphs
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        new_text = self.replace_placeholders_in_text(original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
                
                # Process footer tables (for the footer data table)
                for table_idx, table in enumerate(section.footer.tables):
                    print(f"  Processing footer table {table_idx + 1}")
                    # Adjust table width for margins
                    self.adjust_table_width(table)
                    
                    # Apply blue shading to appropriate footer table rows
                    for row_idx, row in enumerate(table.rows):
                        row_text = ' '.join([cell.text.strip() for cell in row.cells])
                        print(f"    Footer row {row_idx}: '{row_text[:50]}...'")
                        
                        # Apply blue shading to the first row of footer tables (typically headers)
                        if row_idx == 0:
                            print(f"    Applying blue shading to footer table header row {row_idx}")
                            for cell in row.cells:
                                self.add_blue_shading_to_cell(cell)
                    
                    # Replace placeholders in footer table cells
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph.text.strip():
                                    original_text = paragraph.text
                                    new_text = self.replace_placeholders_in_text(original_text)
                                    if new_text != original_text:
                                        paragraph.text = new_text
                                        print(f"  Replaced: '{original_text}' → '{new_text}'")
    
    def add_bullet_points_to_text(self, paragraph, items):
        """Add bullet points to a paragraph for the given items."""
        if not items:
            return
        
        paragraph.clear()
        for item in items:
            run = paragraph.add_run(f"• {item}")
            run.add_break()  # Add line break after each item

    def add_checkbox_points_to_text(self, paragraph, items):
        """Add checkbox bullet points to a paragraph for the given items."""
        if not items:
            return
        
        paragraph.clear()
        for item in items:
            run = paragraph.add_run(f"☐ {item}")  # Unicode checkbox character
            run.add_break()  # Add line break after each item
    
    def process_placeholders_with_formatting(self, paragraph, placeholders_to_check=None):
        """Process paragraph text with special formatting for bullets and checkboxes."""
        if not paragraph.text.strip():
            return False
    
        if placeholders_to_check is None:
            placeholders_to_check = [
                ('{PRE_REQ}', self.data.get('PRE_REQS', []), self.add_bullet_points_to_text),
                ('{REQ_EQUIP}', self.data.get('EQUIPMENT_LIST', []), self.add_bullet_points_to_text),
                ('{COMP_CRIT}', self.data.get('COMPLETION_CRITERIA', []), self.add_checkbox_points_to_text)
            ]
        
        original_text = paragraph.text
        modified = False
        
        for placeholder, items, formatter_func in placeholders_to_check:
            if placeholder in original_text:
                print(f"  Adding formatted items for {placeholder}")
                formatter_func(paragraph, items)
                modified = True
                break
        
        # If no special formatting was applied, just do regular placeholder replacement
        if not modified:
            new_text = self.replace_placeholders_in_text(original_text)
            if new_text != original_text:
                paragraph.text = new_text
                modified = True
    
        return modified
    
    # def add_default_footer(self, footer):
    #     """Add default footer content with placeholders - DISABLED to preserve template footer format."""
    #     pass
    
    def create_fallback_document(self):
        """Create a comprehensive fallback document with all formatting and sections."""
        print("Creating enhanced fallback document...")
        doc = Document()
        self.set_document_margins(doc)

        # Main Title and Info
        mtl_number = self.data.get('MTL_NUMBER', 'MTL 1')
        title_text = 'MASTER TASKLIST: TEACHBACK' if 'MTL 3' in mtl_number else 'MASTER TASK LIST'
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # MTL Title
        mtl_title_para = doc.add_paragraph()
        mtl_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        mtl_title_run = mtl_title_para.add_run(self.data.get('MTL_TITLE', ''))
        mtl_title_run.font.bold = True
        mtl_title_run.font.size = Pt(14)

        # Category and Estimated Time
        info_line = []
        if self.data.get('CATEGORY'):
            info_line.append(f"Category: {self.data['CATEGORY']}")
        if self.data.get('ESTIMATED_TIME'):
            info_line.append(f"Estimated Time: {self.data['ESTIMATED_TIME']}")
        if info_line:
            info_para = doc.add_paragraph(" | ".join(info_line))
            info_para.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()  # Space

        # Main Steps Table
        steps = self.data.get('STEPS', [])
        if steps:
            main_table = doc.add_table(rows=3, cols=3)
            main_table.style = 'Table Grid'
            main_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            self.adjust_table_width(main_table)
            col_row = main_table.rows[0]
            col_row.cells[0].text = "#"
            col_row.cells[1].text = "STEP"
            col_row.cells[2].text = "TECH. INITIALS"
            for cell in col_row.cells:
                self.add_blue_shading_to_cell(cell)
            header_row = main_table.rows[1]
            header_row.cells[0].text = ""
            step_cell = header_row.cells[1]
            step_cell.text = ""
            step_para = step_cell.paragraphs[0]
            step_run = step_para.add_run("STEP")
            step_run.font.bold = True
            detail_para = step_cell.add_paragraph()
            detail_text = ("(Provide limited detail for each step without being too wordy. "
                         "This is a checklist for a trained, experienced employee. "
                         "For more detail, see the MTL2 for this task. Avoid inserting "
                         "photos or other attachments on MTL1.)")
            detail_run = detail_para.add_run(detail_text)
            detail_run.font.size = Pt(8)
            detail_run.font.italic = True
            header_row.cells[2].text = ""
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not run.font.italic:
                            run.font.bold = True
            main_table._element.remove(main_table.rows[2]._element)
            for i, step in enumerate(steps, start=1):
                row = main_table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = step
                row.cells[2].text = ""

        # Prerequisites
        pre_reqs = self.data.get('PREREQUISITES', self.data.get('PRE_REQS', []))
        if pre_reqs:
            doc.add_paragraph()
            prereq_header = doc.add_paragraph()
            prereq_header_run = prereq_header.add_run("PREREQUISITES:")
            prereq_header_run.font.bold = True
            prereq_header_run.font.size = Pt(11)
            for req in pre_reqs:
                req_para = doc.add_paragraph(f"• {req}")
                req_para.paragraph_format.left_indent = Inches(0.25)

        # Safety Notes
        safety_notes = self.data.get('SAFETY_NOTES', [])
        if safety_notes:
            doc.add_paragraph()
            safety_header = doc.add_paragraph()
            safety_header_run = safety_header.add_run("SAFETY NOTES:")
            safety_header_run.font.bold = True
            safety_header_run.font.size = Pt(11)
            for note in safety_notes:
                note_para = doc.add_paragraph(f"• {note}")
                note_para.paragraph_format.left_indent = Inches(0.25)

        # Equipment List
        equipment_list = self.data.get('EQUIPMENT_LIST', [])
        if equipment_list:
            doc.add_paragraph()
            equip_header = doc.add_paragraph()
            equip_header_run = equip_header.add_run("EQUIPMENT LIST:")
            equip_header_run.font.bold = True
            equip_header_run.font.size = Pt(11)
            for equip in equipment_list:
                equip_para = doc.add_paragraph(f"• {equip}")
                equip_para.paragraph_format.left_indent = Inches(0.25)

        # Completion Criteria
        completion_criteria = self.data.get('COMPLETION_CRITERIA', [])
        if completion_criteria:
            doc.add_paragraph()
            criteria_header = doc.add_paragraph()
            criteria_header_run = criteria_header.add_run("COMPLETION CRITERIA:")
            criteria_header_run.font.bold = True
            criteria_header_run.font.size = Pt(11)
            for criteria in completion_criteria:
                criteria_para = doc.add_paragraph(f"☐ {criteria}")
                criteria_para.paragraph_format.left_indent = Inches(0.25)

        # Troubleshooting (as a table)
        troubleshooting = self.data.get('TROUBLESHOOTING', {})
        if troubleshooting:
            doc.add_paragraph()
            trouble_header = doc.add_paragraph()
            trouble_header_run = trouble_header.add_run("TROUBLESHOOTING:")
            trouble_header_run.font.bold = True
            trouble_header_run.font.size = Pt(11)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.rows[0].cells[0].text = "Issue"
            table.rows[0].cells[1].text = "Solution"
            for issue, solution in troubleshooting.items():
                row = table.add_row()
                row.cells[0].text = issue
                row.cells[1].text = solution

        # Related Procedures
        related = self.data.get('RELATED_PROCEDURES', [])
        if related:
            doc.add_paragraph()
            related_header = doc.add_paragraph()
            related_header_run = related_header.add_run("RELATED PROCEDURES:")
            related_header_run.font.bold = True
            related_header_run.font.size = Pt(11)
            for proc in related:
                proc_para = doc.add_paragraph(f"• {proc}")
                proc_para.paragraph_format.left_indent = Inches(0.25)

        # Add enhanced footer section for MTL 2 & 3 (notes and signatures on 2nd page)
        self.add_enhanced_footer_section(doc)
        return doc
    
    def generate_document(self, template_path=None, output_path=None):
        """Generate the final document with full formatting."""
        print(f"\n{'='*60}")
        print(f"FINAL MTL DOCUMENT GENERATION")
        print(f"{'='*60}")
        print(f"Data: {self.data.get('MTL_NUMBER', self.data.get('MTL_#', 'Unknown'))} - {self.data.get('MTL_TITLE', 'Unknown')}")
        
        try:
            if template_path and os.path.exists(template_path):
                print(f"Loading template: {os.path.basename(template_path)}")
                doc = Document(template_path)
                self.process_template_document(doc)
            else:
                if template_path:
                    print(f"⚠️  Template not found: {template_path}")
                print("📄 Creating enhanced fallback document")
                doc = self.create_fallback_document()
        except Exception as e:
            print(f"❌ Error processing template: {e}")
            print("📄 Creating fallback document instead")
            doc = self.create_fallback_document()
        
        # Determine output path
        if not output_path:
            mtl_number = self.data.get('MTL_NUMBER', self.data.get('MTL_#', 'MTL'))
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_dir = "generated_mtls"
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            output_path = os.path.join(output_dir, f"{mtl_number}_Final_{timestamp}.docx")
        
        # Save document
        try:
            doc.save(output_path)
            print(f"\n✅ Final document generated successfully!")
            print(f"📁 Output: {output_path}")
            print(f"✨ Features included:")
            
            # Dynamic header message
            mtl_number = self.data.get('MTL_NUMBER', 'MTL 1')
            if 'MTL 3' in mtl_number:
                print(f"   - MASTER TASKLIST: TEACHBACK header")
            else:
                print(f"   - MASTER TASK LIST header")
                
            print(f"   - Left-justified MTL_TITLE")
            print(f"   - Blue column headers (#00699b)")
            print(f"   - 0.5-inch margins")
            print(f"   - Table width adjusted for margins (7.5\")")
            print(f"   - Enhanced STEP description")
            print(f"   - Template footer format preserved")
            
            # Show specific features for MTL 2
            if 'MTL 2' in mtl_number:
                print(f"   - Tools Required section")
                print(f"   - Completion Criteria section")
                print(f"   - Trainer notes section")
                print(f"   - Complete signature block (trainer & learner)")
                
            # Show trainer notes for MTL 3
            elif 'MTL 3' in mtl_number:
                print(f"   - Trainer notes section")
                print(f"   - Complete signature block (trainer & learner)")
                
            print(f"   - {len(self.data.get('STEPS', []))} steps populated")
            print(f"{'='*60}\n")
            return output_path
        except Exception as e:
            raise Exception(f"Error saving document: {e}")
    
    def replace_text_with_formatting(self, paragraph, original_text, new_text):
        """Replace text in paragraph and apply bold formatting if it's MTL_TITLE."""
        if '{MTL_TITLE}' in original_text:
            # Clear the paragraph and add new text with bold formatting
            paragraph.clear()
            run = paragraph.add_run(new_text)
            run.bold = True
        else:
            # Normal text replacement
            paragraph.text = new_text
    
    def add_enhanced_footer_section(self, doc):
        """Add enhanced footer section with trainer notes and signature block for MTL 2 & 3."""
        mtl_number = self.data.get('MTL_NUMBER', 'MTL 1')
        
        # Add Pre Reqs and Tools Required sections for MTL 2
        if 'MTL 2' in mtl_number or 'MTL2' in mtl_number:
            # Add Pre Reqs section with bullet points
            pre_reqs = self.data.get('PRE_REQS', [])
            if pre_reqs:
                doc.add_paragraph()
                prereq_header = doc.add_paragraph()
                prereq_header_run = prereq_header.add_run("PRE REQS:")
                prereq_header_run.font.bold = True
                prereq_header_run.font.size = Pt(11)
                for req in pre_reqs:
                    req_para = doc.add_paragraph()
                    req_para.paragraph_format.left_indent = Inches(0.25)
                    req_run = req_para.add_run(f"• {req}")
                    req_run.font.size = Pt(10)

            # Add Tools Required section with bullet points
            tools_required = self.data.get('EQUIPMENT_LIST', [])
            if tools_required:
                doc.add_paragraph()
                tools_header = doc.add_paragraph()
                tools_header_run = tools_header.add_run("TOOLS REQUIRED:")
                tools_header_run.font.bold = True
                tools_header_run.font.size = Pt(11)
                
                for tool in tools_required:
                    tool_para = doc.add_paragraph()
                    tool_para.paragraph_format.left_indent = Inches(0.25)
                    tool_run = tool_para.add_run(f"• {tool}")
                    tool_run.font.size = Pt(10)
                    
            # Add Completion Criteria section with checkboxes
            completion_criteria = self.data.get('COMPLETION_CRITERIA', [])
            if completion_criteria:
                doc.add_paragraph()
                criteria_header = doc.add_paragraph()
                criteria_header_run = criteria_header.add_run("COMPLETION CRITERIA:")
                criteria_header_run.font.bold = True
                criteria_header_run.font.size = Pt(11)
                
                for criteria in completion_criteria:
                    criteria_para = doc.add_paragraph()
                    criteria_para.paragraph_format.left_indent = Inches(0.25)
                    criteria_run = criteria_para.add_run(f"☐ {criteria}")  # Unicode checkbox
                    criteria_run.font.size = Pt(10)
    
        # Add trainer notes and signature block to second page for MTL 2 & 3
        if 'MTL 2' in mtl_number or 'MTL2' in mtl_number or 'MTL 3' in mtl_number or 'MTL3' in mtl_number:
            doc.add_page_break()
            
            # Create trainer notes section
            trainer_header = doc.add_paragraph()
            trainer_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
            trainer_run = trainer_header.add_run("TRAINER NOTES:")
            trainer_run.font.bold = True
            trainer_run.font.size = Pt(11)
            
            # Add lines for trainer notes
            for i in range(4):
                line_para = doc.add_paragraph()
                line_para.add_run("_" * 80)
                line_para.paragraph_format.space_after = Pt(6)
            
            # Add comprehensive signature block
            doc.add_paragraph()  # Space before signatures
            
            # Trainer signature section
            trainer_sig_para = doc.add_paragraph()
            trainer_sig_run = trainer_sig_para.add_run("TRAINER SIGNATURE & PRINTED NAME:")
            trainer_sig_run.font.bold = True
            
            trainer_line_para = doc.add_paragraph()
            trainer_line_para.add_run("Signature: _________________________________     Date: ____________")
            
            trainer_name_para = doc.add_paragraph()
            trainer_name_para.add_run("Printed Name: _________________________________")
            
            # Add space between trainer and learner signatures
            doc.add_paragraph()
            
            # Learner signature section
            learner_sig_para = doc.add_paragraph()
            learner_sig_run = learner_sig_para.add_run("LEARNER SIGNATURE & PRINTED NAME:")
            learner_sig_run.font.bold = True
            
            learner_line_para = doc.add_paragraph()
            learner_line_para.add_run("Signature: _________________________________     Date: ____________")
            
            learner_name_para = doc.add_paragraph()
            learner_name_para.add_run("Printed Name: _________________________________")
            
            # Add explanatory text
            doc.add_paragraph()
            explanation_para = doc.add_paragraph()
            explanation_text = ("By signing above, the trainer certifies that the learner has demonstrated "
                              "competency in performing this task according to company standards. "
                              "The learner acknowledges understanding and ability to perform this task safely.")
            explanation_run = explanation_para.add_run(explanation_text)
            explanation_run.font.size = Pt(9)
            explanation_run.font.italic = True
    
    def generate_mtl(self, output_file):
        """Generate a simple DOCX file as a placeholder. Replace with your real logic."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("MTL Document generated by FinalMTLGenerator.")
        # You can use self.data to fill in the document as needed.
        doc.save(output_file)


def main():
    """Main function to run the final generator."""
    import sys
    
    # Check for command line argument
    if len(sys.argv) > 1:
        json_file = sys.argv[1]
    else:
        json_file = "mtl1_data.json"
    
    try:
        # Check if JSON file exists
        if not os.path.exists(json_file):
            print(f"❌ Error: JSON file '{json_file}' not found.")
            return 1

        # Load JSON to determine MTL type
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        mtl_number = data.get('MTL_NUMBER', '').upper()
        if 'MTL 1' in mtl_number:
            template_found = 'MTL/templates/MTL1 Template GT-00X-0X.docx'
        elif 'MTL 2' in mtl_number:
            template_found = 'MTL/templates/MTL2 Template GT-00X-0X.docx'
        elif 'MTL 3' in mtl_number:
            template_found = 'MTL/templates/MTL3 Template GT-00X-0X.docx'
        else:
            template_found = 'MTL/templates/MTL1 Template GT-00X-0X.docx'  # Default fallback

        if not os.path.exists(template_found):
            print(f"⚠️  Template not found: {template_found}")
            template_found = None

        # Create final generator
        generator = FinalMTLGenerator(json_file)

        # Generate document
        output_file = generator.generate_document(template_path=template_found)

        print(f"🎉 SUCCESS! Your final MTL document has been generated.")
        print(f"📂 File location: {os.path.abspath(output_file)}")

    except Exception as e:
        print(f"❌ Error: {e}")
        return 1

    return 0


if __name__ == "__main__":
    exit(main())
