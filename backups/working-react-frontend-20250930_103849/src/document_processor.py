"""
Handles confidential document processing with privacy-first approach
"""


import os
import sys
import json
import click
import logging
from pathlib import Path
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, Any, Optional, List
import PyPDF2
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
import requests


# Data class for processed document
@dataclass
class ProcessedDocument:
    """Container for processed document data"""
    filename: str
    content: str
    metadata: Dict[str, Any]
    chunks: List[str]
    embeddings: Optional[List[List[float]]] = None
    processed_at: Optional[str] = None

# Main document processing class
class DocumentProcessor:
    def __init__(self, ollama_url: str = "http://ollama:11434"):
        self.ollama_url = ollama_url
        self.embedding_model = None
        self.setup_logging()

    def setup_logging(self):
        """Configure logging for document processing"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger("DocumentProcessor")

    def extract_mtl_json(self, processed_doc: ProcessedDocument, mtl_type: str = "MTL 2") -> dict:
        """Generate MTL JSON using LLM for a processed document."""
        prompt = f"""Create a detailed Master Task List (MTL) JSON for casino slot machine technicians. Extract specific technical details from the provided document.

DOCUMENT: {processed_doc.filename}
CONTENT: {processed_doc.content[:2500]}

Create comprehensive JSON with detailed, specific content (not generic placeholders):

{{
  "MTL_NUMBER": "{mtl_type}",
  "MTL_TITLE": "[Extract exact title from document]",
  "VERSION_NUMBER": "GT-002",
  "REVISION_NUMBER": "001",
  "CREATED_DATE": "07/2025",
  "CREATED_BY": "EG Tech Development Team",
  "CATEGORY": "[Specific equipment type from document]",
  "ESTIMATED_TIME": "[Realistic time from document context]",
  "PREREQUISITES": [
    "[List specific certifications mentioned]",
    "[Reference any prerequisite procedures]",
    "[Required training or experience]"
  ],
  "SAFETY_NOTES": [
    "[Extract specific safety procedures from document]",
    "[List lockout/tagout requirements]",
    "[Mention PPE requirements]"
  ],
  "REQUIRED_TOOLS": [
    "[List exact tools mentioned in document]",
    "[Include part numbers if provided]",
    "[Specify software versions mentioned]"
  ],
  "STEPS": [
    "[Create 8-15 detailed steps based on document sections]",
    "[Include specific procedures, not generic ones]",
    "[Mention exact error codes, part numbers, procedures from source]",
    "[Add sub-steps for complex procedures]"
  ],
  "COMPLETION_CRITERIA": [
    "[Specific testing mentioned in document]",
    "[Performance standards from source]",
    "[Documentation requirements]"
  ],
  "TROUBLESHOOTING": {{
    "[Extract specific issues from document]": "[Detailed solutions from source]",
    "[List error codes mentioned]": "[Specific resolution steps]",
    "[Common problems noted]": "[Exact solutions provided]"
  }},
  "RELATED_PROCEDURES": [
    "[Reference other MTL numbers if applicable]",
    "[Mention manufacturer documentation]"
  ],
  "EQUIPMENT_MODELS": [
    "[Extract specific models mentioned in document]",
    "[Include manufacturer names and part numbers]"
  ]
}}

Extract ALL specific technical details, part numbers, error codes, software versions, and procedures from the source document. Create professional-grade content that technicians can follow step-by-step. Only return valid JSON."""
        try:
            # Health check with shorter timeout
            health = requests.get(f"{self.ollama_url}/api/tags", timeout=5)
            if health.status_code != 200 or 'llama3' not in health.text:
                return {"error": "llama3 model not loaded in Ollama"}
            
            response = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": "llama3",
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.4,
                        "num_predict": 3000,
                        "top_p": 0.9,
                        "repeat_penalty": 1.1
                    }
                },
                timeout=180  # Balanced timeout
            )
            if response.status_code == 200:
                result = response.json().get('response', '')
                try:
                    json_start = result.find('{')
                    json_end = result.rfind('}') + 1
                    json_str = result[json_start:json_end]
                    mtl_json = json.loads(json_str)
                    def clean_json(obj):
                        if isinstance(obj, dict):
                            return {k: clean_json(v) for k, v in obj.items() if not k.endswith('_HELP')}
                        elif isinstance(obj, list):
                            return [clean_json(i) for i in obj]
                        else:
                            return obj
                    mtl_json = clean_json(mtl_json)
                    return mtl_json
                except Exception as e:
                    logging.error(f"Error parsing MTL JSON: {e}\nRaw response: {result}")
                    return {"error": "Could not parse MTL JSON", "raw": result}
            else:
                return {"error": f"Ollama API error: {response.status_code}"}
        except Exception as e:
            logging.error(f"Error extracting MTL JSON: {e}")
            return {"error": str(e)}
    
    def load_embedding_model(self):
        """Load local embedding model for document vectorization"""
        if self.embedding_model is None:
            try:
                # Use a lightweight model that runs locally
                self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
                self.logger.info("Loaded local embedding model")
            except Exception as e:
                self.logger.warning(f"Could not load embedding model: {e}")
                self.embedding_model = None
    
    # (Removed incomplete/duplicate extract_text_from_pdf method)
    
    def extract_text_from_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text and metadata from PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                metadata = {
                    "pages": len(pdf_reader.pages),
                    "title": pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else ''
                }
            return text.strip(), metadata
        except Exception as e:
            logging.error(f"Error processing PDF {file_path}: {e}")
            return "", {}
    
    def extract_text_from_docx(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text and metadata from DOCX"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "title": ""
            }
            return text.strip(), metadata
        except Exception as e:
            self.logger.error(f"Error processing DOCX {file_path}: {e}")
            return "", {}
    
    def extract_text_from_txt(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            metadata = {
                "file_size": file_path.stat().st_size,
                "lines": len(text.split('\n'))
            }
            return text.strip(), metadata
        except Exception as e:
            self.logger.error(f"Error processing TXT {file_path}: {e}")
            return "", {}
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks for better processing"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to end at a sentence boundary
            if end < len(text):
                last_period = text.rfind('.', start, end)
                last_newline = text.rfind('\n', start, end)
                boundary = max(last_period, last_newline)
                
                if boundary > start:
                    end = boundary + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks
    
    def generate_embeddings(self, chunks: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks"""
        if self.embedding_model is None:
            self.load_embedding_model()
        
        if self.embedding_model is None:
            return []
        
        try:
            embeddings = self.embedding_model.encode(chunks)
            return embeddings.tolist()
        except Exception as e:
            self.logger.error(f"Error generating embeddings: {e}")
            return []
    
    def process_document(self, file_path: str) -> ProcessedDocument:
        """Process a document and return structured data"""
        path_obj = Path(file_path)
        
        if not path_obj.exists():
            raise FileNotFoundError(f"File not found: {path_obj}")
        
        self.logger.info(f"Processing document: {path_obj.name}")
        
        # Extract text based on file type
        file_extension = path_obj.suffix.lower()
        
        if file_extension == '.pdf':
            text, metadata = self.extract_text_from_pdf(path_obj)
        elif file_extension == '.docx':
            text, metadata = self.extract_text_from_docx(path_obj)
        elif file_extension in ['.txt', '.md']:
            text, metadata = self.extract_text_from_txt(path_obj)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        if not text:
            raise ValueError(f"No text extracted from {path_obj}")
        
        # Chunk the text
        chunks = self.chunk_text(text)
        
        # Generate embeddings
        embeddings = self.generate_embeddings(chunks)
        
        # Add processing metadata
        metadata.update({
            "file_type": file_extension,
            "file_path": str(path_obj),
            "chunks_count": len(chunks),
            "character_count": len(text),
            "word_count": len(text.split())
        })
        
        return ProcessedDocument(
            filename=path_obj.name,
            content=text,
            metadata=metadata,
            chunks=chunks,
            embeddings=embeddings
        )
    
    def analyze_document(self, processed_doc: ProcessedDocument, analysis_type: str = "summary") -> str:
        """Analyze document using Ollama"""
        try:
            if analysis_type == "summary":
                prompt = f"""Please provide a comprehensive summary of the following document:

Title: {processed_doc.filename}
Content: {processed_doc.content[:2000]}...

Provide a summary that includes:
1. Main topics and themes
2. Key findings or points
3. Overall purpose and conclusion
4. Any important details or recommendations

Summary:"""
            
            elif analysis_type == "breakdown":
                prompt = f"""Please provide a detailed breakdown and analysis of the following document:

Title: {processed_doc.filename}
Content: {processed_doc.content[:2000]}...

Provide a breakdown that includes:
1. Document structure and organization
2. Main sections and their purposes
3. Key data points, statistics, or metrics
4. Critical insights and implications
5. Potential action items or next steps

Analysis:"""
            
            elif analysis_type == "questions":
                prompt = f"""Based on the following document, generate important questions that someone should ask:

Title: {processed_doc.filename}
Content: {processed_doc.content[:2000]}...

Generate 10 thoughtful questions that would help someone better understand:
- The document's implications
- Missing information that should be investigated
- Next steps or decisions to be made
- Potential risks or opportunities

Questions:"""
            
            else:
                prompt = f"Please analyze this document: {processed_doc.content[:2000]}..."
            
            # Send request to Ollama
            response = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": "llama3",
                    "prompt": prompt,
                    "stream": False,
                    "options": {"use_mmap": False}
                },
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                return result.get('response', 'No analysis generated')
            else:
                return f"Error: Could not analyze document (Status: {response.status_code})"
                
        except Exception as e:
            self.logger.error(f"Error analyzing document: {e}")
            return f"Error analyzing document: {str(e)}"

@click.command()
@click.option('--file', '-f', required=True, help='Path to the document to process')
@click.option('--analysis', '-a', default='summary', 
              type=click.Choice(['summary', 'breakdown', 'questions']),
              help='Type of analysis to perform')
@click.option('--output', '-o', help='Output file for results (optional)')
@click.option('--verbose', '-v', is_flag=True, help='Verbose output')
def main(file: str, analysis: str, output: Optional[str], verbose: bool):
    """Process and analyze confidential documents locally"""
    
    if verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    try:
        processor = DocumentProcessor()
        
        # Process the document
        click.echo(f"📄 Processing document: {file}")
        processed_doc = processor.process_document(file)
        
        click.echo(f"✅ Document processed successfully!")
        click.echo(f"   - {processed_doc.metadata['character_count']} characters")
        click.echo(f"   - {processed_doc.metadata['word_count']} words")
        click.echo(f"   - {processed_doc.metadata['chunks_count']} chunks")
        
        # Analyze the document
        click.echo(f"\n🤖 Performing {analysis} analysis...")
        analysis_result = processor.analyze_document(processed_doc, analysis)
        
        # Output results
        result = {
            "document": {
                "filename": processed_doc.filename,
                "metadata": processed_doc.metadata
            },
            "analysis_type": analysis,
            "analysis": analysis_result,
            "processed_at": processed_doc.processed_at
        }
        
        if output:
            with open(output, 'w') as f:
                json.dump(result, f, indent=2)
            click.echo(f"📁 Results saved to: {output}")
        else:
            click.echo(f"\n📊 {analysis.title()} Analysis:")
            click.echo("=" * 50)
            click.echo(analysis_result)
            click.echo("=" * 50)
        
    except Exception as e:
        click.echo(f"❌ Error: {str(e)}", err=True)
        sys.exit(1)

if __name__ == "__main__":
    main()
